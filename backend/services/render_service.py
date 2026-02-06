import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SECTION_ORDER = ["Summary", "Skills", "Experience", "Education"]


def _parse_sections(tailored_text: str) -> dict:
    text = tailored_text.strip()

    pattern = r"(?im)^(summary|skills|experience|education)\s*:?\s*$"
    matches = list(re.finditer(pattern, text))

    sections = {k: "" for k in SECTION_ORDER}
    if not matches:
        sections["Experience"] = text
        return sections

    for i, m in enumerate(matches):
        name = m.group(1).capitalize()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections[name] = text[start:end].strip()

    return sections


def _set_base_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_spacer(doc: Document, lines: int = 1):
    for _ in range(lines):
        doc.add_paragraph("")


def _add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)


def _is_probable_role_header(line: str) -> bool:
    """
    Heuristic for lines like:
    Amazon — Business Intelligence Engineer | 2023–2025
    Company | Role | Dates
    """
    l = line.strip()
    if not l:
        return False
    # Contains separators commonly used in role headers
    return ("|" in l) or ("—" in l) or (" - " in l)


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*[\-\*\u2022]\s+", line))


def _clean_bullet(line: str) -> str:
    return re.sub(r"^\s*[\-\*\u2022]\s+", "", line.strip())


def _add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def _render_summary(doc: Document, body: str):
    _add_section_heading(doc, "Summary")
    for ln in [x.strip() for x in body.splitlines() if x.strip()]:
        doc.add_paragraph(ln)
    _add_spacer(doc, 1)


def _render_skills(doc: Document, body: str):
    _add_section_heading(doc, "Skills")

    # If body already has grouped lines, keep them.
    # If it's a long paragraph, we still write as a paragraph for ATS.
    lines = [x.strip() for x in body.splitlines() if x.strip()]
    if len(lines) <= 1:
        doc.add_paragraph(body.strip())
    else:
        for ln in lines:
            doc.add_paragraph(ln)
    _add_spacer(doc, 1)


def _render_experience(doc: Document, body: str):
    _add_section_heading(doc, "Experience")

    lines = [x.rstrip() for x in body.splitlines()]
    for ln in lines:
        if not ln.strip():
            continue

        if _is_bullet(ln):
            _add_bullet(doc, _clean_bullet(ln))
            continue

        # Non-bullet line → treat as a header (company/role/date) or subheader
        p = doc.add_paragraph(ln.strip())
        if _is_probable_role_header(ln):
            for r in p.runs:
                r.bold = True

    _add_spacer(doc, 1)


def _render_education(doc: Document, body: str):
    _add_section_heading(doc, "Education")
    for ln in [x.strip() for x in body.splitlines() if x.strip()]:
        doc.add_paragraph(ln)
    # no trailing spacer required


def render_docx(out_path: str, name: str, title_line: str, contact_line: str, tailored_text: str):
    doc = Document()
    _set_base_style(doc)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(name.strip())
    r.bold = True
    r.font.size = Pt(16)

    if title_line.strip():
        p = doc.add_paragraph(title_line.strip())
        p.runs[0].italic = True

    if contact_line.strip():
        doc.add_paragraph(contact_line.strip())

    _add_spacer(doc, 1)

    sections = _parse_sections(tailored_text)

    if sections.get("Summary", "").strip():
        _render_summary(doc, sections["Summary"])

    if sections.get("Skills", "").strip():
        _render_skills(doc, sections["Skills"])

    if sections.get("Experience", "").strip():
        _render_experience(doc, sections["Experience"])

    if sections.get("Education", "").strip():
        _render_education(doc, sections["Education"])

    doc.save(out_path)
